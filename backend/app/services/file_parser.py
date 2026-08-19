import io
from typing import Tuple
import PyPDF2
import pdfplumber
from docx import Document


class UnsupportedFileTypeError(Exception):
    """Raised when file type is not supported."""

    def __init__(self, filename: str):
        self.filename = filename
        super().__init__(f"Unsupported file type: {filename}")


def detect_file_type(filename: str, mime_type: str = None) -> str:
    """Detect file type from filename extension or MIME type."""
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return "pdf"
    elif filename_lower.endswith(".docx"):
        return "docx"
    elif filename_lower.endswith(".doc"):
        return "doc"  # Old Word format - not fully supported

    # Fallback to MIME type
    if mime_type:
        if "pdf" in mime_type:
            return "pdf"
        elif (
            "wordprocessingml" in mime_type
            or "msword" in mime_type
            or "document" in mime_type
        ):
            return "docx"

    raise UnsupportedFileTypeError(filename)


def parse_pdf(file_bytes: bytes) -> str:
    """Parse PDF file and extract text content."""
    text_parts = []

    # Try pdfplumber first (better for complex layouts)
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception:
        # Fallback to PyPDF2
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    if not text_parts:
        raise ValueError("Could not extract any text from PDF")

    return "\n\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """Parse DOCX file and extract text content."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        if not text_parts:
            raise ValueError("Could not extract any text from DOCX")

        return "\n\n".join(text_parts)

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_cv(file_bytes: bytes, filename: str, mime_type: str = None) -> Tuple[str, str]:
    """
    Parse CV file and extract text content.

    Args:
        file_bytes: Raw file bytes
        filename: Original filename
        mime_type: Optional MIME type

    Returns:
        Tuple of (extracted_text, file_type)
    """
    file_type = detect_file_type(filename, mime_type)

    if file_type == "pdf":
        text = parse_pdf(file_bytes)
    elif file_type == "docx":
        text = parse_docx(file_bytes)
    elif file_type == "doc":
        raise UnsupportedFileTypeError(
            f"{filename} - Old .doc format not supported. Please convert to .docx"
        )
    else:
        raise UnsupportedFileTypeError(filename)

    return text, file_type


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive whitespace
    lines = text.split("\n")
    cleaned_lines = []

    for line in lines:
        # Strip whitespace
        line = line.strip()
        # Skip empty lines or lines with only special characters
        if line and not all(c in "=-_*#" for c in line):
            cleaned_lines.append(line)

    # Join with single newlines, removing excessive blank lines
    result = "\n".join(cleaned_lines)

    # Replace multiple spaces with single space
    while "  " in result:
        result = result.replace("  ", " ")

    return result
